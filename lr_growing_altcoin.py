import pandas as pd
import os
import utils
import math
import sys
import matplotlib.pyplot as plt
from structures import *
from docx.shared import Inches
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.run import WD_BREAK

'''
Constants
'''
TEN_DAYS = 16*24*60*60

class BreakIt(Exception): pass

'''
- given ticker entries and interval duration, splits tickers entries in intervals each lasting 'interval_duration'
- returns list of Interval entries
- duration of the interval is specified in number of seconds
'''
def generate_lending_intervals(duration, entries):
	start_date = float('inf')
	end_date = -float('inf')
	for entry in entries:
		if float(entry.timestamp) < start_date:
			start_date = entry.timestamp
		if float(entry.timestamp) > end_date:
			end_date = entry.timestamp
	if start_date != float('inf') and end_date != -float('inf'):
		num_intervals = int(math.ceil((end_date - start_date)/duration))
		intervals = list()
		for i in range(1, num_intervals):
			period_start = start_date + (i - 1) * duration
			period_end = start_date + i * duration
			if period_end > end_date:
				period_end = end_date
			entries_in_period = list(filter(lambda x: x.timestamp >= period_start and x.timestamp <= period_end, entries))
			interval = LendingInterval(entry.ticker, period_start, period_end, entries_in_period)
			intervals.append(interval)
		return intervals
	else:
		return list()

'''
Returns list of Intervals of interest when lending rate was higher than average 
'''
def get_interest_intervals(lending_intervals):
	interest_intervals = list()
	filtered_interest_intervals = list()
	filteredout_interest_intervals = list()
	for i in range(len(lending_intervals)):
		interest_interval_start = None
		interest_interval_end = None
		lending_tickers = list()
		curr_interval = lending_intervals[i]
		avg_lending_rate = curr_interval.get_avg_lending_rate()
		for entry in curr_interval.lending_entries:
			#begin interest interval when lending rate is higher than average
			if entry.lending_rate >= avg_lending_rate:
				if not interest_interval_start:
					interest_interval_start = entry.timestamp
			else:
				#finish interest interval when lending rate gets below average
				if interest_interval_start:
					interest_interval_end = entry.timestamp
			# collect ticker entries within the period of interest
			if interest_interval_start and not interest_interval_end:
				lending_tickers.append(entry)
			if interest_interval_start and interest_interval_end:
				# ensure that short intervals with less than 10 ticker entries aren't counted
				if len(lending_tickers) > 10:
					interest_interval = InterestInterval(entry.ticker, interest_interval_start, interest_interval_end, lending_tickers)
					interest_intervals.append(interest_interval)
				interest_interval_start = None
				interest_interval_end = None
				lending_tickers = list()
		try:
			#if iteration over entries of curr_interval finished, but lending rate is above average,
			# continue on the next period
			if interest_interval_start and not interest_interval_end:
				num_intervals = 1
				while(num_intervals+i<len(lending_intervals)):
					curr_interval = lending_intervals[i+num_intervals]
					#store previous lending interval
					prev_interval = lending_intervals[i+num_intervals-1]
					for j in range(0, len(curr_interval.lending_entries)):
						#from new interval, we consider at the moment only first j tickers
						tickers_considered = curr_interval.lending_entries[:(j+1)]
						entry = tickers_considered[j]
						#calculate sum of lending rates for entries considered in new period
						lend_rate_new_interval = sum(list(map(lambda x: x.lending_rate, tickers_considered)))
						#average lending rate is now calculated across all tickets considered
						avg_lending_rate = (prev_interval.get_avg_lending_rate()*len(prev_interval.lending_entries)+lend_rate_new_interval)/float(len(prev_interval.lending_entries) + len(tickers_considered))
						#stop interest interval if lending rate is becoming lower than average lending rate over all considered entries
						if entry.lending_rate < avg_lending_rate:
							# ensure that short intervals with less than 10 ticker entries aren't counted
							if len(lending_tickers) > 10:
								interest_interval_end = entry.timestamp
								interest_interval = InterestInterval(entry.ticker, interest_interval_start, interest_interval_end, lending_tickers)
								interest_intervals.append(interest_interval)
							raise BreakIt
						else:
							lending_tickers.append(entry)
					num_intervals += 1
		except BreakIt:
			pass
	for interval in interest_intervals:
		if interval.is_growing():
			filtered_interest_intervals.append(interval)
		else:
			filteredout_interest_intervals.append(interval)
	return filtered_interest_intervals, filteredout_interest_intervals

def main():
	# define period that we are interested in
	period_start = 1480530600.0
	period_end = 1507062600.0

	# collect data about tickers for respective interest period
	btc_rows = utils.get_ticker_data("data/(2016-08-13)-btc_lending_rates_bitfinex.csv", period_start, period_end, 0)
	btc_entries = utils.df_rows_to_lending_entries("BTC", btc_rows, 0, 3)
	print("%d %s entries collected!" % (len(btc_entries), "BTC"))
	xmr_rows = utils.get_ticker_data("data/ltc_bitfinex_data.csv", period_start, period_end, 0)
	xmr_entries = utils.df_rows_to_interest_entries("LTC", xmr_rows, 0, 2, 5)
	print("%d %s entries collected!" % (len(xmr_entries), "TLC"))

	#break data about lending rates into 10 day intervals
	lending_intervals = generate_lending_intervals(TEN_DAYS, btc_entries)

	# return interest intervals that have passed filter function and those which didn't
	filtered_interest_intervals, filteredout_interest_intervals = get_interest_intervals(lending_intervals)
	print("Total filtered interest intervals: %d" % (len(filtered_interest_intervals)))
	print("Total filtered out intereste intervals: %d" % (len(filteredout_interest_intervals)))

	for interval in filtered_interest_intervals:
		entries = list(filter(lambda x: x.timestamp >= interval.start_date and x.timestamp <= interval.end_date, xmr_entries))
		interval.interest_entries = entries
		print("Interval: " + interval.to_string())
	print("------------------------------------")
	print("Total filtered out interest intervals: %d" % (len(filteredout_interest_intervals)))
	for filtered_interval in filteredout_interest_intervals:
		entries = list(filter(lambda x: x.timestamp >= interval.start_date and x.timestamp <= interval.end_date, xmr_entries))
		filtered_interval.interest_entries = entries
		print("Interval: " + filtered_interval.to_string())
	print("------------------------------------")

	# argument stands for the name of Word file
	# TODO: refactor into proper method
	if sys.argv[1]:
		print("Generating Word document...")
		filename = sys.argv[1]
		doc = Document()
		doc.add_heading("LTC Price Analysis based on BTC lending rate (from 2016-08-13 to 2017-09-13), 16 days interval", 0)
		heading = doc.add_heading("Filtered Interest Intervals")
		heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		for idx, filtered_interval in enumerate(filtered_interest_intervals):
			picture_prefix = "ltc_btc_" + str(idx+1)
			utils.plot_interval("LTC", "BTC", filtered_interval, picture_prefix)
			picture_1 = picture_prefix + "_1.png"
			picture_2 = picture_prefix + "_2.png"
			doc.add_picture(picture_1, width=Inches(4.5))
			last_paragraph = doc.paragraphs[-1]
			last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
			doc.add_picture(picture_2, width=Inches(4.5))
			last_paragraph = doc.paragraphs[-1]
			last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
			doc.add_page_break()
		heading = doc.add_heading("Filtered out Interest Intervals")
		heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		for idx, unfiltered_interval in enumerate(filteredout_interest_intervals):
			picture_prefix = "ltc_btc_unfiltered" + str(idx+1)
			utils.plot_interval("LTC", "BTC", filtered_interval, picture_prefix)
			picture_1 = picture_prefix + "_1.png"
			picture_2 = picture_prefix + "_2.png"
			doc.add_picture(picture_1, width=Inches(4.5))
			last_paragraph = doc.paragraphs[-1]
			last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
			doc.add_picture(picture_2, width=Inches(4.5))
			last_paragraph = doc.paragraphs[-1]
			last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
			if idx != len(filteredout_interest_intervals)-1:
				doc.add_page_break()
		doc.save(filename)
		print("Cleaning up...")
		dir_name = os.path.dirname(os.path.realpath(__file__))
		files = os.listdir(dir_name)
		for file in files:
			if file.endswith(".png"):
				os.remove(os.path.join(dir_name, file))

if __name__ == "__main__":
	main()